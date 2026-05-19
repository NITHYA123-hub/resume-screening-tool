# ============================================================
# parser.py — Resume Text Extraction & Entity Parsing
# ============================================================

import re
import os
import io
import logging

# PDF parsing
try:
    import pdfplumber
    PDFPLUMBER_OK = True
except ImportError:
    PDFPLUMBER_OK = False

try:
    import PyPDF2
    PYPDF2_OK = True
except ImportError:
    PYPDF2_OK = False

# DOCX parsing
try:
    from docx import Document
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Known Skills Dictionary ──────────────────────────────
SKILLS_DB = {
    "programming": [
        "python","java","javascript","typescript","c++","c#","ruby","go","rust",
        "kotlin","swift","scala","r","matlab","php","perl","bash","shell","dart",
    ],
    "web": [
        "html","css","react","angular","vue","nodejs","express","django","flask",
        "fastapi","spring","springboot","bootstrap","tailwind","webpack","nextjs",
        "nuxtjs","gatsby","graphql","rest api","soap",
    ],
    "data": [
        "pandas","numpy","scipy","matplotlib","seaborn","plotly","bokeh",
        "tensorflow","pytorch","keras","scikit-learn","xgboost","lightgbm",
        "spark","hadoop","hive","kafka","airflow","dbt","sql","mysql",
        "postgresql","mongodb","redis","elasticsearch","power bi","tableau",
        "excel","looker","databricks","snowflake",
    ],
    "cloud": [
        "aws","azure","gcp","docker","kubernetes","terraform","jenkins",
        "github actions","circleci","ansible","linux","git","github","gitlab",
        "bitbucket","nginx","apache",
    ],
    "ai_ml": [
        "machine learning","deep learning","natural language processing","nlp",
        "computer vision","reinforcement learning","llm","gpt","bert","transformers",
        "data science","feature engineering","model deployment","mlops",
    ],
    "soft": [
        "leadership","communication","teamwork","problem solving","critical thinking",
        "project management","agile","scrum","kanban","jira","confluence",
    ],
}

ALL_SKILLS = []
for cat_skills in SKILLS_DB.values():
    ALL_SKILLS.extend(cat_skills)

# ─── Education Keywords ───────────────────────────────────
EDUCATION_KEYWORDS = [
    "bachelor", "master", "phd", "b.tech", "m.tech", "b.e", "m.e",
    "bsc", "msc", "bca", "mca", "mba", "b.com", "m.com", "diploma",
    "associate", "high school", "secondary", "university", "college",
    "institute", "degree",
]

EDUCATION_PATTERNS = [
    r"(b\.?tech|m\.?tech|b\.?e|m\.?e|b\.?sc|m\.?sc|bca|mca|mba|phd|bachelor|master|diploma)"
    r"[\s\S]{0,80}?(computer science|information technology|electronics|mechanical|civil|"
    r"electrical|data science|artificial intelligence|software|physics|mathematics|statistics)?",
]

# ─── Text Extraction ──────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from PDF bytes."""
    text = ""
    # Try pdfplumber first (better for structured PDFs)
    if PDFPLUMBER_OK:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

    # Fallback to PyPDF2
    if PYPDF2_OK:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from DOCX bytes."""
    if not DOCX_OK:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode plain text file."""
    for enc in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(enc)
        except Exception:
            continue
    return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    else:
        # Try text decode as fallback
        return extract_text_from_txt(file_bytes)


# ─── Entity Extraction ───────────────────────────────────
def extract_name(text: str) -> str:
    """
    Heuristic: the candidate name is usually in the first 3 non-empty lines.
    We look for a line that is:
      - 2-5 words
      - All words capitalized (title case)
      - No digits / common header words
    """
    IGNORE = {"resume", "curriculum", "vitae", "cv", "page", "profile", "summary",
               "contact", "information", "details", "objective", "career"}
    lines = [l.strip() for l in text.split("\n") if l.strip()][:8]
    for line in lines:
        words = line.split()
        if 2 <= len(words) <= 5:
            if all(w[0].isupper() for w in words if w.isalpha()):
                if not any(w.lower() in IGNORE for w in words):
                    if not any(char.isdigit() for char in line):
                        return line
    return "Not Found"


def extract_email(text: str) -> str:
    """Extract first valid email address."""
    pattern = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
    matches = re.findall(pattern, text)
    return matches[0] if matches else "Not Found"


def extract_phone(text: str) -> str:
    """Extract first phone number (various formats)."""
    patterns = [
        r"(\+?\d{1,3}[\s\-]?)?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}",
        r"\+?\d[\d\s\-().]{9,15}\d",
    ]
    for pat in patterns:
        matches = re.findall(pat, text)
        if matches:
            phone = matches[0] if isinstance(matches[0], str) else "".join(matches[0])
            phone = re.sub(r"\s+", " ", phone).strip()
            if len(re.sub(r"\D", "", phone)) >= 10:
                return phone
    return "Not Found"


def extract_skills(text: str) -> list:
    """Match skills from the known skills database."""
    text_lower = text.lower()
    found = []
    for skill in ALL_SKILLS:
        # Use word-boundary aware matching
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found.append(skill.title())
    return list(dict.fromkeys(found))  # preserve order, deduplicate


def extract_education(text: str) -> list:
    """Extract education qualifications."""
    found = []
    text_lower = text.lower()
    lines = text.split("\n")

    for line in lines:
        line_lower = line.lower()
        if any(kw in line_lower for kw in EDUCATION_KEYWORDS):
            cleaned = line.strip()
            if cleaned and len(cleaned) > 5:
                found.append(cleaned)

    # Deduplicate and limit to top 5
    seen = set()
    unique = []
    for item in found:
        key = item.lower()[:40]
        if key not in seen:
            seen.add(key)
            unique.append(item)

    return unique[:5]


def extract_experience_years(text: str) -> str:
    """Try to detect years of experience."""
    patterns = [
        r"(\d+)\+?\s*years?\s+of\s+experience",
        r"experience\s+of\s+(\d+)\+?\s*years?",
        r"(\d+)\+?\s*years?\s+experience",
        r"(\d+)\+?\s*yrs?\.?\s+(?:of\s+)?experience",
    ]
    for pat in patterns:
        match = re.search(pat, text, re.IGNORECASE)
        if match:
            return match.group(1) + " years"
    return "Not specified"


def extract_linkedin(text: str) -> str:
    """Extract LinkedIn profile URL."""
    match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    return match.group(0) if match else "Not Found"


def extract_github(text: str) -> str:
    """Extract GitHub profile URL."""
    match = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    return match.group(0) if match else "Not Found"


# ─── Master Parser ────────────────────────────────────────
def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Full pipeline: extract text → parse all entities.
    Returns a structured dict with all candidate info.
    """
    text = extract_text(file_bytes, filename)
    if not text.strip():
        return {
            "filename": filename,
            "raw_text": "",
            "error": "Could not extract text from file.",
            "name": "Unknown",
            "email": "Not Found",
            "phone": "Not Found",
            "skills": [],
            "education": [],
            "experience": "Not specified",
            "linkedin": "Not Found",
            "github": "Not Found",
        }

    return {
        "filename":   filename,
        "raw_text":   text,
        "error":      None,
        "name":       extract_name(text),
        "email":      extract_email(text),
        "phone":      extract_phone(text),
        "skills":     extract_skills(text),
        "education":  extract_education(text),
        "experience": extract_experience_years(text),
        "linkedin":   extract_linkedin(text),
        "github":     extract_github(text),
    }


# ─── ATS Score Calculation ────────────────────────────────
def compute_ats_score(parsed: dict, job_keywords: list = None) -> dict:
    """
    Compute an ATS compatibility score (0-100) based on:
    - Contact info completeness   (20 pts)
    - Skills detected             (30 pts)
    - Education info              (20 pts)
    - Keyword match (job desc)    (30 pts)
    """
    score = 0
    breakdown = {}

    # Contact completeness
    contact_score = 0
    if parsed.get("name") not in ["Not Found", "Unknown"]:    contact_score += 7
    if parsed.get("email") != "Not Found":                     contact_score += 7
    if parsed.get("phone") != "Not Found":                     contact_score += 6
    score += contact_score
    breakdown["Contact Info"] = contact_score

    # Skills
    skill_count = len(parsed.get("skills", []))
    skill_score = min(30, skill_count * 3)
    score += skill_score
    breakdown["Skills Detected"] = skill_score

    # Education
    edu_score = min(20, len(parsed.get("education", [])) * 7)
    score += edu_score
    breakdown["Education"] = edu_score

    # Keyword match
    if job_keywords:
        raw = parsed.get("raw_text", "").lower()
        matched = sum(1 for kw in job_keywords if kw.lower() in raw)
        kw_score = min(30, int((matched / max(len(job_keywords), 1)) * 30))
    else:
        kw_score = 15  # default mid score when no JD provided
    score += kw_score
    breakdown["Job Match"] = kw_score

    score = min(100, score)
    grade = "Excellent" if score >= 80 else "Good" if score >= 60 else "Average" if score >= 40 else "Poor"

    return {
        "total": score,
        "grade": grade,
        "breakdown": breakdown,
    }


if __name__ == "__main__":
    # Quick smoke-test with a dummy text
    dummy = b"""
    John Smith
    john.smith@email.com | +1 (555) 123-4567 | linkedin.com/in/johnsmith | github.com/johnsmith

    EDUCATION
    B.Tech in Computer Science - MIT University (2018-2022)

    SKILLS
    Python, Machine Learning, TensorFlow, Flask, SQL, Docker, AWS, Git

    EXPERIENCE
    3 years of experience as a Data Scientist at TechCorp.
    """
    result = parse_resume(dummy, "test_resume.txt")
    for k, v in result.items():
        if k != "raw_text":
            print(f"{k:12}: {v}")
    ats = compute_ats_score(result, ["python", "machine learning", "sql"])
    print(f"\nATS Score: {ats['total']}/100 - {ats['grade']}")
    print("Breakdown:", ats["breakdown"])
